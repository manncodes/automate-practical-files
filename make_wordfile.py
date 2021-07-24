from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Practical ', 0)

p = document.add_paragraph('')
p.add_run('Aim:').bold = True



p = document.add_paragraph('')
p.add_run('Program code:').bold = True


p = document.add_paragraph('')
p.add_run('Screenshots').bold = True


p = document.add_paragraph('')
p.add_run('Conclusions').bold = True


document.add_page_break()

document.save('practical 1.docx')